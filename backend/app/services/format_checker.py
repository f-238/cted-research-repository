import json
from pathlib import Path
from docx import Document
from pypdf import PdfReader

REQUIRED_SECTIONS = [
    "Introduction",
    "Methodology",
    "Results and Discussion",
    "Conclusion",
    "References",
]


def _docx_text(document: Document) -> str:
    return "\n".join(p.text for p in document.paragraphs)


def check_document(path: str) -> dict:
    suffix = Path(path).suffix.lower()
    if suffix == ".docx":
        return check_docx(path)
    if suffix == ".pdf":
        return check_pdf(path)
    return {"is_compliant": False, "warnings": ["Unsupported file type."], "passed_items": [], "raw_summary": {}}


def check_docx(path: str) -> dict:
    document = Document(path)
    warnings: list[str] = []
    passed: list[str] = []
    summary: dict = {}

    fonts, sizes, alignments = set(), set(), set()
    for paragraph in document.paragraphs:
        if paragraph.alignment is not None:
            alignments.add(str(paragraph.alignment))
        for run in paragraph.runs:
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.size:
                sizes.add(round(run.font.size.pt, 1))

    text = _docx_text(document).lower()
    page_count = getattr(document.core_properties, "pages", None)
    if page_count:
        summary["page_count"] = page_count
        if page_count > 12:
            warnings.append("More than 12 pages.")
    missing_sections = []
    for section in REQUIRED_SECTIONS:
        aliases = [section.lower()]
        if section == "Methodology":
            aliases.append("methods")
        if not any(alias in text for alias in aliases):
            missing_sections.append(section)
    if missing_sections:
        warnings.append("Missing IMRAD section: " + ", ".join(missing_sections))
    else:
        passed.append("IMRAD sections are present")

    has_tnr = any("times new roman" in font.lower() for font in fonts)
    has_arial = any("arial" in font.lower() for font in fonts)
    valid_size = (has_tnr and 12.0 in sizes) or (has_arial and 11.0 in sizes)
    if fonts and not (has_tnr or has_arial):
        warnings.append("Invalid font. Use Times New Roman or Arial.")
    else:
        passed.append("Font style appears compliant or inherited from style")
    if sizes and not valid_size:
        warnings.append("Invalid font size. Use Times New Roman 12 pt or Arial 11 pt.")
    else:
        passed.append("Font size appears compliant or inherited from style")

    section = document.sections[0]
    margins = {
        "top": round(section.top_margin.inches, 2),
        "bottom": round(section.bottom_margin.inches, 2),
        "left": round(section.left_margin.inches, 2),
        "right": round(section.right_margin.inches, 2),
    }
    summary["margins_inches"] = margins
    if any(value < 0.9 or value > 1.25 for value in margins.values()):
        warnings.append("Margins should be close to 1 inch on all sides.")
    else:
        passed.append("Margins appear compliant")

    page = (round(section.page_width.inches, 2), round(section.page_height.inches, 2))
    summary["page_size_inches"] = page
    if page not in [(8.5, 11.0), (8.27, 11.69)]:
        warnings.append("Invalid paper size. Use Short/Letter or A4.")
    else:
        passed.append("Page size appears compliant")

    spacing_values = [p.paragraph_format.line_spacing for p in document.paragraphs if p.paragraph_format.line_spacing]
    summary["line_spacing_samples"] = [str(value) for value in spacing_values[:10]]
    if not spacing_values:
        warnings.append("Line spacing is not double or could not be fully detected.")
    elif not any(str(value) in {"2.0", "DOUBLE (2)"} or value == 2 for value in spacing_values):
        warnings.append("Line spacing is not double.")
    else:
        passed.append("Line spacing appears compliant")

    heading_hits = [p.text for p in document.paragraphs if p.style and "Heading" in p.style.name]
    summary["heading_count"] = len(heading_hits)
    if len(heading_hits) < 3:
        warnings.append("Heading styles appear incomplete; use template heading styles.")
    else:
        passed.append("Heading styles detected")

    if not alignments:
        warnings.append("Paragraph alignment could not be detected; verify justified/body alignment.")
    else:
        passed.append("Paragraph alignment metadata detected")

    summary["fonts"] = sorted(fonts)
    summary["font_sizes"] = sorted(sizes)
    summary["compliance_status"] = "Passed" if not warnings else "Warning"
    return {"is_compliant": not warnings, "warnings": warnings, "passed_items": passed, "raw_summary": summary}


def check_pdf(path: str) -> dict:
    reader = PdfReader(path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages).lower()
    missing_sections = []
    for section in REQUIRED_SECTIONS:
        aliases = [section.lower()]
        if section == "Methodology":
            aliases.append("methods")
        if not any(alias in text for alias in aliases):
            missing_sections.append(section)
    warnings = []
    passed = []
    page_count = len(reader.pages)
    if page_count > 12:
        warnings.append("More than 12 pages.")
    if missing_sections:
        warnings.append("Missing IMRAD section: " + ", ".join(missing_sections))
    else:
        passed.append("IMRAD sections are present")
    for page in reader.pages:
        width = round(float(page.mediabox.width) / 72, 2)
        height = round(float(page.mediabox.height) / 72, 2)
        normalized = tuple(sorted((width, height)))
        if normalized not in [tuple(sorted((8.5, 11.0))), tuple(sorted((8.27, 11.69)))]:
            warnings.append("Invalid paper size. Use Short/Letter or A4.")
            break
    warnings.append("PDF format checking is limited. Upload DOCX for accurate font, font size, and line spacing validation.")
    return {
        "is_compliant": not warnings,
        "warnings": warnings,
        "passed_items": passed,
        "raw_summary": {"page_count": page_count, "compliance_status": "Passed" if not warnings else "Warning"},
    }


def serialize_check(result: dict) -> tuple[bool, str, str, str]:
    return (
        result["is_compliant"],
        json.dumps(result["warnings"]),
        json.dumps(result["passed_items"]),
        json.dumps(result["raw_summary"]),
    )
